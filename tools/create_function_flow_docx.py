from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "luong-chinh-cac-chuc-nang.md"
OUTPUT_DOCX = Path(os.environ.get("FUNCTION_FLOW_DOCX", r"C:\Users\TUNA\Downloads\luong_chinh_cac_chuc_nang.docx"))


def set_run_font(run, size: float = 11, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc: Document, text: str, style: str | None = None, bold: bool = False, size: float = 11):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def add_heading(doc: Document, text: str, level: int) -> None:
    style = "Title" if level == 1 else f"Heading {min(level, 3)}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(8 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True)


def shade_cell(cell, fill: str = "EAF2F8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(clean_inline(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, cell_text in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[i], cell_text, bold=True, size=10)
        shade_cell(table.rows[0].cells[i])
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            set_cell_text(cells[i], cell_text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(set(cell.replace(":", "").replace("-", "").strip()) == set() for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def main() -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    in_code_block = False
    code_lines: list[str] = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            if in_code_block:
                add_para(doc, "\n".join(code_lines), style="Intense Quote", size=9.5)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("|"):
            table, i = parse_md_table(lines, i)
            add_table(doc, table)
            continue
        if stripped.startswith("# "):
            add_heading(doc, clean_inline(stripped[2:].strip()), 1)
        elif stripped.startswith("## "):
            add_heading(doc, clean_inline(stripped[3:].strip()), 2)
        elif stripped.startswith("### "):
            add_heading(doc, clean_inline(stripped[4:].strip()), 3)
        elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            add_para(doc, stripped.strip("*"), bold=True)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(clean_inline(stripped[2:].strip()))
            set_run_font(run)
        elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. ") or stripped.startswith("4. ") or stripped.startswith("5. ") or stripped.startswith("6. ") or stripped.startswith("7. ") or stripped.startswith("8. ") or stripped.startswith("9. "):
            add_para(doc, clean_inline(stripped))
        else:
            add_para(doc, clean_inline(stripped))
        i += 1

    try:
        doc.save(OUTPUT_DOCX)
        print(OUTPUT_DOCX)
    except PermissionError:
        fallback = OUTPUT_DOCX.with_name(f"{OUTPUT_DOCX.stem}_v2{OUTPUT_DOCX.suffix}")
        doc.save(fallback)
        print(fallback)


if __name__ == "__main__":
    main()
