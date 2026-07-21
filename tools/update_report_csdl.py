from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = Path(r"C:\Users\TUNA\Downloads\bao_cao_source_for_codex.docx")
OUTPUT_DOCX = Path(r"C:\Users\TUNA\Downloads\bao_cao_da_sua_CSDL.docx")
SOURCE_MD = ROOT / "docs" / "mo-ta-co-so-du-lieu.md"


def set_run_font(run, size: float = 11, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str = "EAF2F8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_para(doc: Document, text: str, style: str | None = None, bold: bool = False, size: float = 11):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=13 if level == 2 else 12, bold=True)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    widths = (2200, 2300, 4700) if cols == 3 else (3000, 6200)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, text in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[i], text, bold=True, size=10)
        shade_cell(table.rows[0].cells[i])
        if i < len(widths):
            set_cell_width(table.rows[0].cells[i], widths[i])
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            if i < len(widths):
                set_cell_width(cells[i], widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if not all(set(cell.replace(":", "").replace("-", "").strip()) == set() for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def strip_intro(md: str) -> list[str]:
    lines = md.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("Hệ thống được thiết kế"))
    return lines[start:]


def append_markdown_section(doc: Document, md: str) -> None:
    lines = strip_intro(md)
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            table, i = parse_md_table(lines, i)
            add_table(doc, table)
            continue
        if line.startswith("## "):
            text = line[3:].strip()
            add_heading(doc, text, 2)
        elif line.startswith("### "):
            text = line[4:].strip()
            add_heading(doc, text, 3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:].strip())
            set_run_font(run)
        else:
            add_para(doc, line)
        i += 1


def remove_old_csdl_section(doc: Document) -> None:
    start_el = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("III.") and "CSDL" in text:
            start_el = p._element
            break
    if start_el is None:
        raise RuntimeError("Cannot find section III. He thong CSDL")

    body = doc._body._element
    children = list(body)
    start_index = children.index(start_el)
    sect_pr = children[-1] if children and children[-1].tag == qn("w:sectPr") else None
    for child in children[start_index:]:
        if child is not sect_pr:
            body.remove(child)


def main() -> None:
    doc = Document(SOURCE_DOCX)
    remove_old_csdl_section(doc)
    add_para(doc, "III. Hệ thống CSDL", bold=True, size=14)
    append_markdown_section(doc, SOURCE_MD.read_text(encoding="utf-8"))
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
