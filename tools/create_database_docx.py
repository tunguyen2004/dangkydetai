from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "mo-ta-co-so-du-lieu.md"
OUTPUT_DOCX = Path(
    os.environ.get(
        "DATABASE_DOCX",
        r"C:\Users\TUNA\Downloads\mo_ta_co_so_du_lieu_dang_ky_de_tai.docx",
    )
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F7F9FC"
BORDER = "A6B3C1"


def clean_inline(text: str) -> str:
    return text.replace("`", "").replace("**", "").strip()


def set_run_font(run, size: float = 11, bold: bool = False, color: str | None = None, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int]) -> None:
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=size, bold=bold)


def add_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False, size: float = 11) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    style = f"Heading {min(level, 3)}"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
    paragraph.paragraph_format.space_after = Pt(7 if level <= 2 else 5)
    run = paragraph.add_run(clean_inline(text))
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, size=9.2, name="Consolas")


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        is_separator = all(set(cell.replace(":", "").replace("-", "").strip()) == set() for cell in cells)
        if not is_separator:
            rows.append(cells)
        i += 1
    return rows, i


def table_widths(cols: int) -> list[int]:
    if cols == 2:
        return [3000, 6360]
    if cols == 3:
        return [2100, 2500, 4760]
    if cols == 4:
        return [2100, 2100, 2100, 3060]
    return [int(9360 / cols)] * cols


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_geometry(table, table_widths(cols))

    for idx, text in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[idx], text, bold=True, size=9.3)
        shade_cell(table.rows[0].cells[idx], HEADER_FILL)
    repeat_header_row(table.rows[0])

    for row_values in rows[1:]:
        cells = table.add_row().cells
        for idx in range(cols):
            text = row_values[idx] if idx < len(row_values) else ""
            set_cell_text(cells[idx], text, size=8.7)
            if len(rows) > 8 and len(table.rows) % 2 == 0:
                shade_cell(cells[idx], LIGHT_FILL)

    set_table_geometry(table, table_widths(cols))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Calibri"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Hệ thống đăng ký đề tài - Mô tả cơ sở dữ liệu")
    set_run_font(run, size=9, color="666666")


def main() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("Mô tả cơ sở dữ liệu hệ thống đăng ký đề tài")
    set_run_font(title_run, size=20, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Bản chốt theo mô hình topics + topic_classes, phạm vi một học kỳ")
    set_run_font(subtitle_run, size=11, color="666666")

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(clean_inline(line))
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 1)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2)
        elif stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            run = paragraph.add_run(clean_inline(stripped[2:]))
            set_run_font(run, size=11)
        elif stripped.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(clean_inline(stripped[2:]))
            set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
        elif stripped[:2].isdigit() and stripped[2:4] == ". ":
            add_paragraph(doc, stripped)
        else:
            add_paragraph(doc, stripped)
        i += 1

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT_DOCX)
        print(OUTPUT_DOCX)
    except PermissionError:
        fallback = OUTPUT_DOCX.with_name(f"{OUTPUT_DOCX.stem}_v2{OUTPUT_DOCX.suffix}")
        doc.save(fallback)
        print(fallback)


if __name__ == "__main__":
    main()
