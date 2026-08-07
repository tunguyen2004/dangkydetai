from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'Bao_cao_He_thong_quan_ly_dang_ky_de_tai_Nhom_10.docx'


# standard_business_brief preset, resolved explicitly for a formal report.
INK = '0B2545'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'
MUTED = '5F6B7A'
GREEN = '0E7490'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')

    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, color=INK, size=9.2, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Trang ')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(field)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    p_pr.append(keep)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.add_run(text)
    set_keep_with_next(p)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='Normal')
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        set_cell_shading(header.cells[idx], LIGHT_BLUE)
        set_cell_text(header.cells[idx], value, bold=True, color=DARK_BLUE, size=font_size)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F4F6F9')
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(label + ' ')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_flow_table(doc, title, rows):
    add_heading(doc, title, 2)
    return add_table(
        doc,
        ['Bước', 'Tác nhân', 'Dữ liệu vào', 'Xử lý và kiểm tra', 'Kết quả'],
        rows,
        [560, 1120, 1800, 3620, 2260],
        font_size=8.3,
    )


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

# Styles, all values explicitly set for the chosen design preset.
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 16, 8),
    ('Heading 2', 13, BLUE, 12, 6),
    ('Heading 3', 12, DARK_BLUE, 8, 4),
]:
    style = styles[style_name]
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ['List Bullet', 'List Number']:
    style = styles[style_name]
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.167

# Header/footer for all but title page.
header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_run = header.add_run('HỆ THỐNG QUẢN LÝ ĐĂNG KÝ ĐỀ TÀI')
header_run.font.name = 'Calibri'
header_run.font.size = Pt(8.5)
header_run.font.color.rgb = RGBColor.from_string(MUTED)
footer = section.footer.paragraphs[0]
add_page_number(footer)

# Cover page
for text, size, bold, color, space_after in [
    ('TRƯỜNG ĐẠI HỌC SƯ PHẠM HÀ NỘI 3', 13, True, INK, 3),
    ('MÔN HỌC: CÔNG NGHỆ WEB', 11, True, MUTED, 0),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

doc.add_paragraph().paragraph_format.space_after = Pt(40)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('BÁO CÁO DỰ ÁN CUỐI KỲ')
run.font.name = 'Calibri'
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string(BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
run = p.add_run('HỆ THỐNG QUẢN LÝ ĐĂNG KÝ ĐỀ TÀI')
run.font.name = 'Calibri'
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string(INK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(42)
run = p.add_run('Quản lý lớp học phần, nhóm sinh viên, đề tài và phê duyệt đăng ký')
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string(MUTED)

cover = add_table(
    doc,
    ['Thông tin', 'Nội dung'],
    [
        ['Nhóm thực hiện', 'Nhóm 10'],
        ['Giảng viên hướng dẫn', '[Điền họ tên giảng viên]'],
        ['Học phần', 'Công nghệ Web'],
        ['Năm học', '2025 - 2026'],
    ],
    [2700, 6660],
    font_size=10,
)
for cell in cover.columns[0].cells:
    set_cell_shading(cell, LIGHT_GRAY)
    for run in cell.paragraphs[0].runs:
        run.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(70)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Hà Nội, tháng 8 năm 2026')
run.italic = True
run.font.size = Pt(11)

doc.add_page_break()

# Contents
add_heading(doc, 'MỤC LỤC', 1)
toc_rows = [
    ['1', 'Mô tả sản phẩm'],
    ['2', 'Phạm vi, tác nhân và yêu cầu nghiệp vụ'],
    ['3', 'Kiến trúc và công nghệ sử dụng'],
    ['4', 'Thiết kế cơ sở dữ liệu'],
    ['5', 'Các luồng nghiệp vụ chính'],
    ['6', 'Chức năng đã xây dựng'],
    ['7', 'Bảo mật, kiểm tra dữ liệu và nhật ký hoạt động'],
    ['8', 'Kiểm thử chức năng'],
    ['9', 'Phân công công việc'],
    ['10', 'Hướng dẫn đóng gói và nộp bài'],
]
add_table(doc, ['STT', 'Nội dung'], toc_rows, [900, 8460], font_size=10)
add_callout(doc, 'LƯU Ý:', 'Mục phân công ở phần 9 có các ô [Điền ...]. Nhóm cần thay bằng thông tin và phần trăm đóng góp thực tế trước khi nộp.')
doc.add_page_break()

# 1. Product description
add_heading(doc, '1. Mô tả sản phẩm', 1)
add_heading(doc, '1.1. Bài toán đặt ra', 2)
add_body(doc, 'Trong quá trình tổ chức học phần hoặc bài tập lớn, giảng viên cần theo dõi việc sinh viên lập nhóm và chọn đề tài. Nếu thực hiện bằng tin nhắn, biểu mẫu rời rạc hoặc danh sách thủ công thì khó kiểm soát số lượng thành viên, số nhóm đã chọn một đề tài, thời hạn đăng ký và lịch sử duyệt. Sản phẩm được xây dựng để tập trung hóa các thao tác này trên một hệ thống web.')

add_heading(doc, '1.2. Mục tiêu', 2)
add_bullets(doc, [
    'Quản lý tập trung tài khoản, học phần, lớp học phần, sinh viên và giảng viên phụ trách.',
    'Thiết lập đợt đăng ký có thời gian tạo nhóm và thời gian đăng ký đề tài rõ ràng.',
    'Cho phép sinh viên tạo nhóm, mời thành viên, chuyển quyền trưởng nhóm và theo dõi trạng thái đăng ký.',
    'Cho phép giảng viên chuẩn bị đề tài gốc, mở đề tài cho lớp/đợt phù hợp và duyệt đăng ký có phản hồi.',
    'Đảm bảo ràng buộc nghiệp vụ: đúng quyền, đúng lớp, đúng đợt, đúng thời hạn, không trùng nhóm và không vượt sức chứa đề tài.',
])

add_heading(doc, '1.3. Phạm vi triển khai', 2)
add_body(doc, 'Dự án triển khai trong phạm vi một học kỳ hiện tại. Một học phần có thể có nhiều lớp học phần; một giảng viên có thể phụ trách nhiều lớp; một đợt đăng ký có thể áp dụng cho nhiều lớp. Hệ thống tập trung vào đăng ký đề tài bài tập lớn, chưa triển khai phần quản lý tiến độ, nộp báo cáo, chấm điểm, gợi ý đề tài bằng AI hoặc ghép nhóm tự động.')

add_heading(doc, '1.4. Giá trị của sản phẩm', 2)
add_table(doc, ['Đối tượng', 'Giá trị nhận được'], [
    ['Quản trị viên', 'Chuẩn bị dữ liệu nền, kiểm soát tài khoản, lớp, thời gian đăng ký và truy vết hoạt động.'],
    ['Giảng viên', 'Quản lý kho đề tài, mở đề tài linh hoạt cho lớp/đợt, hỗ trợ tạo nhóm và xử lý đăng ký có phản hồi.'],
    ['Sinh viên', 'Biết rõ thời hạn, tự lập nhóm hoặc tham gia qua lời mời, chỉ trưởng nhóm đăng ký và xem được kết quả xử lý.'],
], [2300, 7060], font_size=9.5)

# 2 actors
add_heading(doc, '2. Phạm vi, tác nhân và yêu cầu nghiệp vụ', 1)
add_heading(doc, '2.1. Tác nhân của hệ thống', 2)
add_body(doc, 'Tác nhân là người sử dụng ở bên ngoài hệ thống: Quản trị viên, Giảng viên và Sinh viên. Hệ thống không được xem là tác nhân; các việc như kiểm tra session, kiểm tra quyền, thời hạn và ràng buộc dữ liệu là phần xử lý nội bộ.')
add_table(doc, ['Vai trò', 'Chức năng chính'], [
    ['Quản trị viên', 'Quản lý tài khoản; khóa/mở; reset mật khẩu; quản lý học phần; tạo/sửa/xóa lớp; gán sinh viên; tạo/gán/mở/đóng đợt đăng ký; xem nhật ký hoạt động.'],
    ['Giảng viên', 'Tạo đề tài gốc; mở/đóng đề tài cho lớp và đợt; xem, tìm kiếm và phân trang nhóm; tạo nhóm hỗ trợ; duyệt, từ chối hoặc hủy duyệt đăng ký.'],
    ['Sinh viên', 'Đăng nhập; tạo nhóm khi được cho phép; mời hoặc phản hồi lời mời; chuyển quyền trưởng nhóm; đăng ký đề tài khi là trưởng nhóm; theo dõi trạng thái và phản hồi.'],
], [1900, 7460], font_size=9.2)

add_heading(doc, '2.2. Quy tắc nghiệp vụ quan trọng', 2)
add_bullets(doc, [
    'Một sinh viên chỉ được thuộc một nhóm trong cùng một lớp học phần và cùng một đợt đăng ký.',
    'Một nhóm phải có đúng một trưởng nhóm; chỉ trưởng nhóm được gửi, hủy đăng ký đề tài hoặc chuyển quyền trưởng nhóm.',
    'Sinh viên chỉ được tạo nhóm khi lớp cho phép tự tạo và đang trong khoảng thời gian tạo nhóm của đợt đăng ký mở.',
    'Khi nhận lời mời, hệ thống kiểm tra lại lớp, số thành viên tối đa và việc sinh viên đã thuộc nhóm khác hay chưa.',
    'Một nhóm chỉ có một đăng ký đề tài còn hiệu lực tại một thời điểm. Trạng thái còn hiệu lực là pending hoặc approved.',
    'Một đề tài được mở cho từng lớp/đợt có giới hạn max_groups riêng; số chỗ được tính theo các đăng ký pending và approved.',
    'Giảng viên chỉ được xem hoặc xử lý nhóm, đề tài và đăng ký thuộc lớp do mình phụ trách.',
])

# Architecture
add_heading(doc, '3. Kiến trúc và công nghệ sử dụng', 1)
add_heading(doc, '3.1. Kiến trúc triển khai', 2)
add_body(doc, 'Ứng dụng được xây dựng theo mô hình web PHP thuần. Trình duyệt gửi yêu cầu đến các trang PHP; phần nghiệp vụ thực hiện kiểm tra session, quyền và dữ liệu trước khi truy cập MySQL qua PDO. Kết quả sau đó được hiển thị bằng HTML, CSS và JavaScript. Cấu trúc này phù hợp với phạm vi môn học, dễ triển khai trên Laragon hoặc máy chủ hỗ trợ PHP/MySQL.')
add_table(doc, ['Tầng', 'Thành phần', 'Vai trò'], [
    ['Giao diện', 'HTML5, CSS3, JavaScript', 'Hiển thị giao diện responsive cho điện thoại, máy tính bảng và máy tính; hỗ trợ popup, bộ lọc, tìm kiếm, phân trang và thông báo.'],
    ['Xử lý ứng dụng', 'PHP 8', 'Xử lý biểu mẫu, phân quyền, session, kiểm tra nghiệp vụ, điều hướng, thông báo và ghi nhật ký.'],
    ['Truy cập dữ liệu', 'PDO + MySQL', 'Thực hiện truy vấn có tham số, quản lý transaction ở các thao tác cần tính nhất quán.'],
    ['Dịch vụ hỗ trợ', 'PHPMailer / SMTP', 'Gửi liên kết đặt lại mật khẩu; token chỉ dùng một lần và có thời hạn.'],
], [1700, 2400, 5260], font_size=9.2)

add_heading(doc, '3.2. Cấu trúc thư mục chính', 2)
add_table(doc, ['Thư mục / file', 'Vai trò'], [
    ['index.php', 'Điểm vào của dự án; chuyển người dùng đã đăng nhập đến dashboard phù hợp và hiển thị trang giới thiệu cho khách.'],
    ['admin/', 'Các trang quản trị: tài khoản, học phần, lớp học phần, đợt đăng ký và nhật ký hoạt động.'],
    ['teacher/', 'Các trang giảng viên: nhóm, đề tài và duyệt đăng ký.'],
    ['student/', 'Các trang sinh viên: nhóm, lời mời, chuyển quyền trưởng nhóm và đăng ký đề tài.'],
    ['auth/', 'Đăng nhập, đăng xuất, đổi mật khẩu, quên mật khẩu và đặt lại mật khẩu.'],
    ['includes/', 'Khởi tạo ứng dụng, kết nối CSDL, helper dùng chung, header và footer.'],
    ['database/', 'Các file SQL khởi tạo và dữ liệu mẫu phục vụ chạy/kiểm thử dự án.'],
    ['assets/', 'Tệp CSS, JavaScript và hình ảnh giao diện.'],
], [2900, 6460], font_size=9.2)

add_callout(doc, 'ĐIỂM VÀO DỰ ÁN:', 'File index.php được đặt ngay tại thư mục mã nguồn. Khi truy cập dự án, hệ thống khởi chạy từ file này theo đúng yêu cầu nộp bài.')

# DB
add_heading(doc, '4. Thiết kế cơ sở dữ liệu', 1)
add_heading(doc, '4.1. Nguyên tắc thiết kế', 2)
add_body(doc, 'CSDL được thiết kế xoay quanh nghiệp vụ học phần - lớp học phần - đợt đăng ký - nhóm sinh viên - đề tài - đăng ký đề tài. Các thực thể có quan hệ nhiều-nhiều được tách qua bảng trung gian để không lặp dữ liệu và để mở rộng được cho nhiều lớp hoặc nhiều đợt.')
add_bullets(doc, [
    'Tách topics và topic_classes: nội dung đề tài gốc không bị lặp khi giảng viên mở cùng đề tài cho nhiều lớp/đợt.',
    'Tách registration_periods và registration_period_classes: một đợt có thể áp dụng cho nhiều lớp, một lớp cũng có thể tham gia nhiều đợt.',
    'Tách student_groups và group_members: không lưu cứng số thành viên; hệ thống đếm thành viên thực tế từ group_members.',
    'Lưu lịch sử đăng ký bằng topic_registrations thay vì sửa đè hoặc xóa bản ghi cũ khi bị từ chối, hủy hoặc hủy duyệt.',
])

add_heading(doc, '4.2. Danh sách bảng dữ liệu', 2)
table_rows = [
    ['users', 'Tài khoản, vai trò, mã người dùng, trạng thái khóa và yêu cầu đổi mật khẩu.'],
    ['password_reset_tokens', 'Token đặt lại mật khẩu đã băm, thời hạn và trạng thái sử dụng.'],
    ['courses', 'Danh mục học phần.'],
    ['classes', 'Lớp học phần, học phần thuộc về, giảng viên phụ trách và quy định nhóm mặc định.'],
    ['class_students', 'Liên kết sinh viên với lớp học phần.'],
    ['registration_periods', 'Thông tin, thời hạn và trạng thái nháp/mở/đóng của đợt đăng ký.'],
    ['registration_period_classes', 'Liên kết đợt đăng ký với lớp được áp dụng.'],
    ['topics', 'Đề tài gốc do giảng viên tạo.'],
    ['topic_classes', 'Mở đề tài cho lớp và đợt cụ thể, có max_groups và trạng thái mở/đóng.'],
    ['student_groups', 'Nhóm sinh viên của một lớp trong một đợt.'],
    ['group_members', 'Thành viên thực tế và vai trò trưởng nhóm/thành viên.'],
    ['group_invitations', 'Lời mời vào nhóm cùng trạng thái phản hồi.'],
    ['topic_registrations', 'Lịch sử nhóm đăng ký đề tài và kết quả xử lý của giảng viên.'],
    ['activity_logs', 'Nhật ký các thao tác nghiệp vụ quan trọng.'],
]
add_table(doc, ['Bảng', 'Mục đích'], table_rows, [3000, 6360], font_size=8.8)

add_heading(doc, '4.3. Quan hệ dữ liệu cốt lõi', 2)
add_table(doc, ['Quan hệ', 'Ý nghĩa'], [
    ['courses 1 - n classes', 'Một học phần có thể có nhiều lớp học phần.'],
    ['classes n - n users qua class_students', 'Một lớp có nhiều sinh viên; sinh viên có thể học nhiều lớp.'],
    ['registration_periods n - n classes qua registration_period_classes', 'Xác định đợt đăng ký áp dụng cho những lớp nào.'],
    ['topics n - n classes/đợt qua topic_classes', 'Một đề tài gốc được tái sử dụng cho nhiều lớp và đợt mà không nhân bản nội dung.'],
    ['student_groups 1 - n group_members', 'Một nhóm có nhiều thành viên; vai trò leader được lưu tại group_members.role.'],
    ['student_groups 1 - n topic_registrations', 'Một nhóm có thể có lịch sử nhiều lần đăng ký nhưng chỉ một bản ghi còn hiệu lực.'],
], [4300, 5060], font_size=9.0)

add_heading(doc, '4.4. Các trạng thái quan trọng', 2)
add_table(doc, ['Đối tượng', 'Trạng thái', 'Ý nghĩa'], [
    ['Đợt đăng ký', 'draft', 'Đợt nháp, chỉ Admin dùng để cấu hình.'],
    ['Đợt đăng ký', 'open', 'Cho phép thao tác theo đúng các mốc thời gian đã thiết lập.'],
    ['Đợt đăng ký', 'closed', 'Ngừng tạo nhóm/đăng ký mới nhưng giữ nguyên lịch sử.'],
    ['Đăng ký đề tài', 'pending / approved', 'Đang giữ ràng buộc với đề tài; nhóm chưa được gửi đăng ký khác.'],
    ['Đăng ký đề tài', 'rejected / cancelled / revoked', 'Không còn hiệu lực; giữ lịch sử để truy vết. Revoked là hủy kết quả đã duyệt.'],
    ['Nhóm', 'forming / registered / locked', 'Đang lập nhóm / đã đăng ký đề tài / bị khóa theo trạng thái nghiệp vụ.'],
], [1900, 2900, 4560], font_size=9.0)

# Flows
add_heading(doc, '5. Các luồng nghiệp vụ chính', 1)
add_body(doc, 'Các luồng sau được mô tả theo nguyên tắc: tác nhân là người dùng; hệ thống thực hiện kiểm tra session, quyền, thời hạn và dữ liệu ở phần xử lý.')

add_flow_table(doc, '5.1. Luồng đăng nhập và xác thực quyền', [
    ['1', 'Người dùng', 'Email, mật khẩu', 'Kiểm tra dữ liệu bắt buộc và định dạng email.', 'Yêu cầu hợp lệ về định dạng.'],
    ['2', 'Người dùng', 'Thông tin đăng nhập', 'Tra cứu users, kiểm tra tài khoản không bị khóa và đối chiếu password_hash.', 'Xác thực thành công hoặc báo lỗi.'],
    ['3', 'Người dùng', 'Tài khoản hợp lệ', 'Tạo lại session ID; lưu user_id, role, last_activity. Nếu must_change_password = 1 thì chuyển đến đổi mật khẩu.', 'Phiên đăng nhập an toàn.'],
    ['4', 'Người dùng', 'URL hoặc thao tác', 'Mỗi trang/POST kiểm tra session, timeout 20 phút, role và phạm vi dữ liệu theo lớp.', 'Cho phép hoặc điều hướng kèm thông báo.'],
])

add_flow_table(doc, '5.2. Luồng Admin chuẩn bị dữ liệu và mở đợt', [
    ['1', 'Admin', 'Tài khoản, học phần, lớp', 'Kiểm tra role admin; tạo tài khoản với user_code tự sinh theo vai trò; tạo học phần, lớp, phân công giảng viên và gán sinh viên.', 'Dữ liệu nền hoàn chỉnh.'],
    ['2', 'Admin', 'Tên đợt và các mốc thời gian', 'Kiểm tra group_start < group_end và register_start < register_end; lưu registration_periods trạng thái draft.', 'Đợt nháp để cấu hình.'],
    ['3', 'Admin', 'Một hoặc nhiều lớp', 'Gán lớp vào registration_period_classes; tránh gán trùng lớp vào cùng một đợt.', 'Xác định đúng phạm vi áp dụng.'],
    ['4', 'Admin', 'Lệnh mở đợt', 'Chỉ cho mở khi có lớp áp dụng và mốc thời gian hợp lệ.', 'Đợt chuyển open; giảng viên/sinh viên trong lớp được phép thao tác.'],
])

add_flow_table(doc, '5.3. Luồng giảng viên tạo đề tài và mở cho lớp', [
    ['1', 'Giảng viên', 'Mã, tên, mô tả, min/max thành viên', 'Kiểm tra giảng viên đăng nhập và mã đề tài không trùng trong phạm vi giảng viên; lưu topics.', 'Có đề tài gốc độc lập.'],
    ['2', 'Giảng viên', 'Đề tài, lớp, đợt, max_groups', 'Chỉ cho chọn lớp do giảng viên phụ trách; tạo topic_classes và kiểm tra lớp đã thuộc đợt.', 'Đề tài được mở đúng lớp/đợt.'],
    ['3', 'Giảng viên', 'Lệnh mở/đóng', 'Cập nhật status mapping; không xóa lịch sử đăng ký đã phát sinh.', 'Kiểm soát khả năng nhận đăng ký theo từng lớp/đợt.'],
])

add_flow_table(doc, '5.4. Luồng tạo nhóm, mời thành viên và chuyển quyền trưởng nhóm', [
    ['1', 'Sinh viên / Giảng viên', 'Lớp, đợt, tên nhóm, trưởng nhóm', 'Kiểm tra lớp/đợt hợp lệ, thời hạn tạo nhóm, số nhóm tối đa và sinh viên chưa ở nhóm khác. Giảng viên có thể tạo nhóm hỗ trợ.', 'Nhóm mới ở trạng thái forming.'],
    ['2', 'Trưởng nhóm', 'Email hoặc mã sinh viên', 'Kiểm tra cùng lớp, lời mời pending không trùng, nhóm còn chỗ và người được mời chưa thuộc nhóm khác.', 'Lời mời pending được tạo.'],
    ['3', 'Sinh viên', 'Chấp nhận / từ chối', 'Khi chấp nhận, kiểm tra lại thời hạn, số người và tình trạng tham gia nhóm trước khi thêm group_members.', 'Thành viên được thêm hoặc lời mời kết thúc.'],
    ['4', 'Trưởng nhóm', 'Một thành viên hiện tại', 'Transaction khóa danh sách thành viên; chuyển role leader/member chỉ khi nhóm không bị khóa và chưa có đăng ký hiệu lực.', 'Luôn còn đúng một trưởng nhóm.'],
])

add_flow_table(doc, '5.5. Luồng đăng ký và duyệt đề tài', [
    ['1', 'Trưởng nhóm', 'topic_class_id', 'Kiểm tra role leader, thời hạn đăng ký, số thành viên thực tế, đề tài mở đúng lớp/đợt và nhóm chưa có đăng ký hiệu lực.', 'Yêu cầu hợp lệ.'],
    ['2', 'Trưởng nhóm', 'Đề tài đã chọn', 'Đếm đăng ký pending/approved để không vượt max_groups; tạo topic_registrations = pending.', 'Đăng ký chờ duyệt.'],
    ['3', 'Giảng viên', 'Đăng ký, phản hồi', 'Kiểm tra giảng viên phụ trách lớp và trạng thái hiện tại. Cập nhật approved hoặc rejected cùng người duyệt/thời điểm duyệt.', 'Kết quả và phản hồi đến nhóm.'],
    ['4', 'Giảng viên', 'Đăng ký approved, lý do', 'Không xóa dữ liệu cũ; chuyển status sang revoked khi cần hủy duyệt.', 'Giữ lịch sử; nhóm có thể đăng ký lại nếu còn thời hạn.'],
])

add_flow_table(doc, '5.6. Luồng đặt lại mật khẩu và nhật ký', [
    ['1', 'Người dùng', 'Email', 'Tạo token ngẫu nhiên, lưu token_hash và hạn dùng; gửi liên kết qua SMTP nếu cấu hình hoàn chỉnh.', 'Liên kết đặt lại mật khẩu một lần.'],
    ['2', 'Người dùng', 'Token, mật khẩu mới', 'Kiểm tra token chưa dùng/chưa hết hạn; mã hóa mật khẩu mới, đánh dấu token used_at và tạo session mới.', 'Đăng nhập bằng mật khẩu mới.'],
    ['3', 'Admin / GV / SV', 'Thao tác quan trọng', 'Sau thao tác, ghi activity_logs gồm user_id, action, detail, created_at.', 'Có dữ liệu truy vết.'],
    ['4', 'Admin', 'Bộ lọc log', 'Kiểm tra role admin; tìm kiếm/lọc activity_logs theo người dùng, thao tác và thời gian.', 'Theo dõi lịch sử nghiệp vụ.'],
])

# Features
add_heading(doc, '6. Chức năng đã xây dựng', 1)
add_heading(doc, '6.1. Phân hệ Quản trị viên', 2)
add_table(doc, ['Nhóm chức năng', 'Nội dung đã thực hiện'], [
    ['Quản lý tài khoản', 'Tạo tài khoản; tự sinh mã AD/GV/SV; kiểm tra trùng email; kiểm tra số điện thoại 10 chữ số; khóa/mở tài khoản; reset mật khẩu và yêu cầu đổi mật khẩu.'],
    ['Học phần và lớp', 'Tạo/sửa/xóa học phần; tạo/sửa/xóa lớp học phần; phân công giảng viên; gán nhiều sinh viên; bỏ sinh viên khỏi lớp khi chưa phát sinh ràng buộc.'],
    ['Đợt đăng ký', 'Tạo/sửa đợt bằng popup; gán nhiều lớp; thay đổi trạng thái nháp/mở/đóng; thiết lập các mốc thời gian.'],
    ['Tra cứu dữ liệu', 'Tìm kiếm, bộ lọc và phân trang cho danh sách lớp; lọc dữ liệu qua query string để không mất trạng thái khi tải lại trang.'],
    ['Nhật ký', 'Trang nhật ký hoạt động có tìm kiếm, bộ lọc và phân trang; chỉ Admin được truy cập.'],
], [2350, 7010], font_size=9.0)

add_heading(doc, '6.2. Phân hệ Giảng viên', 2)
add_table(doc, ['Nhóm chức năng', 'Nội dung đã thực hiện'], [
    ['Nhóm sinh viên', 'Tìm kiếm theo tên nhóm/mã tham gia/sinh viên; lọc theo lớp, đợt và trạng thái; phân trang; tạo nhóm hỗ trợ bằng popup.'],
    ['Đề tài', 'Tạo đề tài gốc; mở đề tài cho lớp/đợt; giới hạn số nhóm; đóng/mở; xóa khi chưa có đăng ký; tìm kiếm/lọc/phân trang.'],
    ['Duyệt đăng ký', 'Lọc theo trạng thái; tìm theo nhóm, đề tài, mã sinh viên, lớp; lọc theo lớp/đợt; duyệt, từ chối và hủy duyệt kèm phản hồi.'],
], [2350, 7010], font_size=9.0)

add_heading(doc, '6.3. Phân hệ Sinh viên', 2)
add_table(doc, ['Nhóm chức năng', 'Nội dung đã thực hiện'], [
    ['Nhóm của bạn', 'Tạo nhóm theo đợt/lớp, xem thành viên, mời qua email hoặc mã người dùng, phản hồi lời mời, hủy lời mời.'],
    ['Trưởng nhóm', 'Chuyển quyền trưởng nhóm cho thành viên hợp lệ bằng transaction; không cho thay đổi khi nhóm đã có đăng ký hiệu lực hoặc bị khóa.'],
    ['Đăng ký đề tài', 'Chỉ trưởng nhóm được gửi; hiển thị đề tài phù hợp lớp/đợt; kiểm tra số lượng thành viên, sức chứa đề tài, thời hạn và trạng thái đăng ký.'],
    ['Theo dõi kết quả', 'Xem trạng thái chờ duyệt, đã duyệt, từ chối, đã hủy, hủy duyệt và phản hồi của giảng viên.'],
], [2350, 7010], font_size=9.0)

# Security
add_heading(doc, '7. Bảo mật, kiểm tra dữ liệu và nhật ký hoạt động', 1)
add_table(doc, ['Nội dung', 'Cách áp dụng trong dự án'], [
    ['Xác thực mật khẩu', 'Mật khẩu được lưu bằng password_hash và kiểm tra bằng password_verify, không lưu mật khẩu dạng văn bản thuần.'],
    ['Phân quyền máy chủ', 'Mỗi trang bảo vệ gọi require_login/require_role; không chỉ ẩn nút ở giao diện. Dữ liệu theo lớp tiếp tục được kiểm tra chủ sở hữu.'],
    ['Session', 'Đổi session ID sau đăng nhập; lưu thời điểm hoạt động gần nhất; tự đăng xuất sau 20 phút không hoạt động.'],
    ['CSRF', 'Biểu mẫu POST có csrf_token và được kiểm tra trước khi thực hiện thao tác thay đổi dữ liệu.'],
    ['SQL injection', 'Sử dụng PDO prepared statements cho dữ liệu do người dùng nhập.'],
    ['XSS', 'Dữ liệu đưa ra HTML được xử lý bằng hàm e() dựa trên htmlspecialchars.'],
    ['Đặt lại mật khẩu', 'Token được lưu dạng băm, có hạn dùng, chỉ sử dụng một lần; sau reset tạo lại session đăng nhập.'],
    ['Ghi log', 'Lưu các thao tác nghiệp vụ quan trọng, không ghi nhận kiểu theo dõi người dùng đang ở trang nào trong bao lâu.'],
], [2500, 6860], font_size=9.0)

# Test plan
add_heading(doc, '8. Kiểm thử chức năng', 1)
add_body(doc, 'Kiểm thử được thực hiện theo các tình huống nghiệp vụ chính và các lỗi nhập liệu thường gặp. Bảng sau là danh sách kiểm thử đại diện; khi nộp bài có thể bổ sung ảnh chụp màn hình minh chứng nếu giảng viên yêu cầu.')
test_rows = [
    ['TC01', 'Đăng nhập đúng/sai', 'Kiểm tra điều hướng theo vai trò và thông báo lỗi xác thực.'],
    ['TC02', 'Truy cập sai quyền', 'Sinh viên truy cập trang Admin hoặc Giảng viên phải bị chuyển hướng kèm thông báo.'],
    ['TC03', 'Tài khoản trùng email', 'Form phải báo lỗi nghiệp vụ, không để lộ SQLSTATE.'],
    ['TC04', 'Số điện thoại sai', 'Chỉ nhận 10 chữ số; không nhận ký tự chữ.'],
    ['TC05', 'max_groups quá lớn', 'UI và backend chặn trước khi ghi CSDL, không để lỗi Numeric value out of range.'],
    ['TC06', 'Tạo nhóm sai thời hạn', 'Chặn thao tác nếu đợt không open hoặc ngoài khung group_start/group_end.'],
    ['TC07', 'Mời vượt số lượng', 'Không tạo/chấp nhận lời mời khi nhóm đã đạt max_members hoặc người được mời đã có nhóm.'],
    ['TC08', 'Đăng ký trùng đề tài', 'Kiểm tra max_groups theo topic_classes và số đăng ký pending/approved.'],
    ['TC09', 'Một nhóm gửi hai đăng ký', 'Ràng buộc active_group_id ngăn nhiều đăng ký còn hiệu lực.'],
    ['TC10', 'Duyệt/hủy duyệt', 'Duyệt chỉ khi pending; hủy duyệt chuyển revoked, không xóa lịch sử.'],
    ['TC11', 'Chuyển quyền trưởng nhóm', 'Transaction đảm bảo luôn có đúng một leader sau khi chuyển.'],
    ['TC12', 'Responsive', 'Các bảng/danh sách chuyển sang dạng cuộn hoặc card phù hợp trên mobile, tablet, laptop.'],
]
add_table(doc, ['Mã', 'Tình huống', 'Kết quả mong đợi'], test_rows, [800, 2580, 5980], font_size=8.8)

# Distribution
add_heading(doc, '9. Phân công công việc', 1)
add_body(doc, 'Bảng sau là mẫu phân công theo vai trò công việc. Nhóm cần thay phần trong ngoặc vuông bằng họ tên, MSSV và tỷ lệ đóng góp thực tế; xóa các hàng không sử dụng để phản ánh đúng số lượng thành viên.')
assignment_rows = [
    ['1', '[Họ tên thành viên 1]', '[MSSV]', 'Phân tích nghiệp vụ, thiết kế CSDL, tích hợp luồng đăng ký đề tài, kiểm thử nghiệp vụ.', '[..]%'],
    ['2', '[Họ tên thành viên 2]', '[MSSV]', 'Phát triển phân hệ Admin: tài khoản, học phần, lớp học phần, đợt đăng ký, nhật ký.', '[..]%'],
    ['3', '[Họ tên thành viên 3]', '[MSSV]', 'Phát triển phân hệ Giảng viên: nhóm, đề tài, duyệt đăng ký, tìm kiếm và phân trang.', '[..]%'],
    ['4', '[Họ tên thành viên 4]', '[MSSV]', 'Phát triển phân hệ Sinh viên: nhóm, lời mời, chuyển trưởng nhóm, đăng ký đề tài.', '[..]%'],
    ['5', '[Họ tên thành viên 5]', '[MSSV]', 'Thiết kế giao diện responsive, kiểm thử, dữ liệu mẫu, triển khai và hoàn thiện báo cáo.', '[..]%'],
]
add_table(doc, ['STT', 'Họ và tên', 'MSSV', 'Công việc phụ trách', 'Tỷ lệ'], assignment_rows, [500, 1900, 1300, 4500, 1160], font_size=8.6)
add_callout(doc, 'YÊU CẦU:', 'Tổng tỷ lệ đóng góp của các thành viên phải bằng 100%. Không nên để tên hoặc MSSV mẫu trong bản nộp chính thức.')

# Submission
add_heading(doc, '10. Hướng dẫn đóng gói và nộp bài', 1)
add_heading(doc, '10.1. Thành phần cần nộp', 2)
add_table(doc, ['Thành phần', 'Yêu cầu'], [
    ['Báo cáo', 'Nộp file Word hoặc PDF. Bản này được xuất ở cả hai định dạng để nhóm lựa chọn.'],
    ['Mã nguồn', 'Nộp đầy đủ thư mục dự án. File index.php nằm tại thư mục gốc và là điểm khởi chạy hệ thống.'],
    ['CSDL', 'Nộp file SQL. Tên CSDL phải theo cú pháp cnw_k3_msv, trong đó msv là MSSV người nộp bài.'],
], [2100, 7260], font_size=9.5)

add_heading(doc, '10.2. Chuẩn bị file SQL trước khi nộp', 2)
add_bullets(doc, [
    'Sao chép file SQL triển khai sang file mới, ví dụ: database/cnw_k3_715105006.sql.',
    'Đổi hai lệnh đầu file SQL thành CREATE DATABASE IF NOT EXISTS `cnw_k3_715105006` ... và USE `cnw_k3_715105006`; thay 715105006 bằng MSSV người nộp bài.',
    'Cập nhật DB_NAME trong file cấu hình local/production tương ứng khi kiểm thử lại file SQL vừa đổi tên.',
    'Không nộp mật khẩu SMTP, mật khẩu CSDL thật hoặc file cấu hình chứa thông tin bí mật lên Git công khai.',
])

add_heading(doc, '10.3. Hướng dẫn chạy nhanh', 2)
add_bullets(doc, [
    'Cài môi trường PHP 8.x, MySQL/MariaDB và web server (có thể dùng Laragon).',
    'Tạo/import CSDL từ file SQL đã đổi tên theo yêu cầu.',
    'Cấu hình DB_HOST, DB_PORT, DB_NAME, DB_USER, DB_PASS trong cấu hình môi trường của máy chạy.',
    'Truy cập thư mục dự án qua web server; hệ thống khởi chạy từ index.php.',
])

# Conclusion
add_heading(doc, '11. Kết luận', 1)
add_body(doc, 'Hệ thống quản lý đăng ký đề tài đã đáp ứng luồng nghiệp vụ cốt lõi: Admin chuẩn bị dữ liệu và mở đợt; Giảng viên quản lý đề tài, nhóm và duyệt; Sinh viên lập nhóm và đăng ký đề tài. Thiết kế dữ liệu tách rõ đề tài gốc với việc mở đề tài cho lớp/đợt, giữ lịch sử đăng ký và kiểm soát các ràng buộc quan trọng. Trong hướng phát triển tiếp theo, hệ thống có thể bổ sung import Excel, thông báo email hoàn chỉnh, xuất báo cáo, ghép nhóm ngẫu nhiên và quản lý tiến độ theo giai đoạn.')

doc.save(OUT)
print(OUT)
